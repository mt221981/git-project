"""File processing service for extracting and cleaning text from documents."""

import hashlib
import re
from pathlib import Path
from typing import Optional
import io
import tempfile
import os

# Document processing libraries
import PyPDF2
from docx import Document

# Windows COM support for .doc files
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False


class FileProcessingError(Exception):
    """Exception raised when file processing fails."""
    pass


class FileProcessor:
    """
    Service for processing uploaded files.

    Handles:
    - Hash calculation for duplicate detection
    - Text extraction from PDF, DOCX, TXT files
    - Text cleaning and normalization
    """

    def __init__(self):
        """Initialize file processor."""
        pass

    def calculate_hash(self, file_content: bytes) -> str:
        """
        Calculate SHA-256 hash of file content.

        Args:
            file_content: Raw file bytes

        Returns:
            Hexadecimal hash string

        Example:
            >>> processor = FileProcessor()
            >>> hash_val = processor.calculate_hash(b"test content")
            >>> len(hash_val)
            64
        """
        sha256_hash = hashlib.sha256()
        sha256_hash.update(file_content)
        return sha256_hash.hexdigest()

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text from file based on file type.

        Args:
            file_content: Raw file bytes
            file_type: File extension (e.g., '.pdf', '.docx', '.txt')

        Returns:
            Extracted text content

        Raises:
            FileProcessingError: If extraction fails or file type is unsupported

        Example:
            >>> processor = FileProcessor()
            >>> text = processor.extract_text(pdf_bytes, '.pdf')
        """
        file_type = file_type.lower()

        extractors = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_doc,  # Old Word format using COM
            '.txt': self._extract_from_txt,
        }

        extractor = extractors.get(file_type)
        if not extractor:
            raise FileProcessingError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {', '.join(extractors.keys())}"
            )

        try:
            text = extractor(file_content)

            if not text or not text.strip():
                raise FileProcessingError("No text could be extracted from file")

            return text

        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from {file_type}: {str(e)}")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file using PyPDF2.

        Args:
            file_content: PDF file bytes

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If PDF extraction fails
        """
        try:
            # Create PDF reader from bytes
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                try:
                    pdf_reader.decrypt('')  # Try empty password
                except:
                    raise FileProcessingError("PDF is password protected")

            # Extract text from all pages
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                if page_text:
                    text_parts.append(page_text)

            if not text_parts:
                raise FileProcessingError("No text found in PDF")

            # Join all pages with double newline
            full_text = '\n\n'.join(text_parts)

            return full_text

        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"PDF extraction failed: {str(e)}")

    def _extract_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file using python-docx.

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If DOCX extraction fails
        """
        try:
            # Create document from bytes
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            if not paragraphs:
                raise FileProcessingError("No text found in DOCX")

            # Join paragraphs with single newline
            full_text = '\n'.join(paragraphs)

            return full_text

        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"DOCX extraction failed: {str(e)}")

    def _extract_from_doc(self, file_content: bytes) -> str:
        """
        Extract text from old .doc files (Word 97-2003) using Windows COM.

        Args:
            file_content: DOC file bytes

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If DOC extraction fails
        """
        if not HAS_WIN32COM:
            raise FileProcessingError(
                "Cannot extract text from .doc files: pywin32 is not installed. "
                "Please save the file as .docx or .pdf format."
            )

        temp_file = None
        word = None

        try:
            # Write content to a temporary file (COM requires file path)
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name

            # Initialize Word application
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open the document
            doc = word.Documents.Open(temp_path)

            # Extract text from the document
            text = doc.Content.Text

            # Close the document
            doc.Close(False)

            if not text or not text.strip():
                raise FileProcessingError("No text found in DOC file")

            return text

        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"DOC extraction failed: {str(e)}")
        finally:
            # Clean up
            if word:
                try:
                    word.Quit()
                except:
                    pass
            if temp_file and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass

    def _extract_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from TXT file with proper encoding handling.

        Tries UTF-8 first, then falls back to other common encodings
        used in Hebrew documents.

        Args:
            file_content: TXT file bytes

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If text extraction fails
        """
        # Try different encodings in order of likelihood
        encodings = [
            'utf-8',
            'utf-8-sig',  # UTF-8 with BOM
            'windows-1255',  # Hebrew Windows encoding
            'iso-8859-8',  # Hebrew ISO encoding
            'cp1255',  # Hebrew code page
            'latin-1',  # Fallback
        ]

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)

                # Check if text looks reasonable (has actual content)
                if text.strip():
                    return text

            except (UnicodeDecodeError, AttributeError):
                continue

        # If all encodings failed
        raise FileProcessingError(
            "Could not decode text file with any supported encoding. "
            "Please ensure the file is in UTF-8, Windows-1255, or ISO-8859-8 format."
        )

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text.

        Performs:
        - Remove page numbers (standalone numbers on lines)
        - Fix broken lines (join lines that were split mid-word)
        - Remove extra whitespace
        - Normalize Hebrew text
        - Remove headers/footers patterns

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text

        Example:
            >>> processor = FileProcessor()
            >>> cleaned = processor.clean_text(raw_text)
        """
        if not text:
            return ""

        # Remove page numbers (lines that are just numbers)
        # Match standalone numbers, optionally with dashes or dots
        text = re.sub(r'^\s*[-–—]?\s*\d+\s*[-–—]?\s*$', '', text, flags=re.MULTILINE)

        # Remove common header/footer patterns
        # Example: "עמוד 5 מתוך 20" or "Page 5 of 20"
        text = re.sub(r'עמוד\s+\d+\s+מתוך\s+\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)

        # Remove repeated header patterns (same line repeated on multiple pages)
        lines = text.split('\n')
        cleaned_lines = []
        seen_headers = {}

        for line in lines:
            stripped = line.strip()

            # Skip empty lines (we'll normalize them later)
            if not stripped:
                cleaned_lines.append('')
                continue

            # Detect potential headers (short lines that repeat)
            if len(stripped) < 100:  # Headers are usually short
                if stripped in seen_headers:
                    seen_headers[stripped] += 1
                    # If line repeats more than 3 times, it's likely a header
                    if seen_headers[stripped] > 3:
                        continue  # Skip this line
                else:
                    seen_headers[stripped] = 1

            cleaned_lines.append(line)

        text = '\n'.join(cleaned_lines)

        # Fix broken lines - join lines that were split mid-word or mid-sentence
        # This is conservative: only join if line doesn't end with sentence-ending punctuation
        lines = text.split('\n')
        fixed_lines = []
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # Check if this line should be joined with next
            if i < len(lines) - 1:
                next_line = lines[i + 1].lstrip()

                # Join if current line doesn't end with sentence-ending punctuation
                # and next line doesn't start with capital letter or number (new section)
                if (line and next_line and
                    not re.search(r'[.!?:;]\s*$', line) and  # No sentence-ending punctuation
                    not re.match(r'^[\d\u0590-\u05FF]', next_line)):  # Not starting new section

                    # Join with space
                    fixed_lines.append(line + ' ' + next_line)
                    i += 2
                    continue

            fixed_lines.append(line)
            i += 1

        text = '\n'.join(fixed_lines)

        # Normalize Hebrew text
        text = self._normalize_hebrew(text)

        # Remove extra whitespace
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)

        # Replace multiple newlines with maximum 2 newlines (paragraph break)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove spaces at start and end of lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        # Final trim
        text = text.strip()

        return text

    def _normalize_hebrew(self, text: str) -> str:
        """
        Normalize Hebrew text.

        - Remove nikud (vowel points)
        - Fix common character issues
        - Normalize quotes

        Args:
            text: Text to normalize

        Returns:
            Normalized text
        """
        if not text:
            return ""

        # Remove Hebrew nikud (vowel points) - Unicode range U+0591 to U+05C7
        text = re.sub(r'[\u0591-\u05C7]', '', text)

        # Normalize quotes - replace various quote types with standard ones
        # Replace fancy quotes with standard quotes
        text = text.replace('"', '"').replace('"', '"')  # English smart quotes
        text = text.replace(''', "'").replace(''', "'")  # English smart apostrophes
        text = text.replace('״', '"').replace('׳', "'")  # Hebrew quotes

        # Normalize dashes - replace various dash types with standard hyphen/en-dash
        text = text.replace('—', '-').replace('–', '-')  # Em dash and en dash to hyphen

        # Remove zero-width characters
        text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)

        # Normalize whitespace around Hebrew punctuation
        # Remove space before Hebrew punctuation marks
        text = re.sub(r'\s+([.,;:!?)])', r'\1', text)
        # Add space after Hebrew punctuation if not already there
        text = re.sub(r'([.,;:!?])([^\s\d])', r'\1 \2', text)

        return text

    def get_text_stats(self, text: str) -> dict:
        """
        Get statistics about extracted text.

        Args:
            text: Text to analyze

        Returns:
            Dictionary with statistics:
            {
                'char_count': int,
                'word_count': int,
                'line_count': int,
                'paragraph_count': int,
                'has_hebrew': bool,
                'has_english': bool
            }
        """
        if not text:
            return {
                'char_count': 0,
                'word_count': 0,
                'line_count': 0,
                'paragraph_count': 0,
                'has_hebrew': False,
                'has_english': False
            }

        # Count characters
        char_count = len(text)

        # Count words (split by whitespace)
        words = text.split()
        word_count = len(words)

        # Count lines
        lines = text.split('\n')
        line_count = len(lines)

        # Count paragraphs (separated by blank lines)
        paragraphs = re.split(r'\n\s*\n', text)
        paragraph_count = len([p for p in paragraphs if p.strip()])

        # Detect languages
        has_hebrew = bool(re.search(r'[\u0590-\u05FF]', text))
        has_english = bool(re.search(r'[a-zA-Z]', text))

        return {
            'char_count': char_count,
            'word_count': word_count,
            'line_count': line_count,
            'paragraph_count': paragraph_count,
            'has_hebrew': has_hebrew,
            'has_english': has_english
        }
